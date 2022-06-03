# -*- encoding:utf-8 -*-
from __future__ import print_function

from docx import Document
from docx.shared import Inches
from opencc import OpenCC
import os
from selenium import webdriver
from bs4 import BeautifulSoup
import requests
import time
import random
import shutil
import sys
import codecs
from textrank4zh import TextRank4Keyword, TextRank4Sentence


# txt轉繁體------------------------------------------------------------------------
def txt_font(output_path_cht, output_path_zh, uuid_sample):
    file_cht = open(output_path_cht + uuid_sample + '.txt', 'r', encoding='utf-8').read()
    converter = OpenCC('s2t')
    file_zh = converter.convert(file_cht)
    f = open(output_path_zh + uuid_sample + '.txt', 'w', encoding='utf-8')
    f.write(file_zh)
    f.close()


# 寫WORD ---------------------------------------------------------
def word(output_path, uuid_sample):
    output_txt = output_path + uuid_sample + '.txt'
    output_word = output_path + uuid_sample + '.docx'

    # 新建WORD
    document = Document()
    # 讀檔
    article_file = open(output_txt, 'r', encoding='utf-8').read()
    # 分割文件
    Article = article_file.split()
    # print(Article)

    # 段落和圖片
    for i in range(len(Article)):
        paragraph = document.add_paragraph(Article[i])
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Inches(0.25)
        document.add_picture(output_path + "image/" + uuid_sample + '_p{}.png'.format(i + 1), width=Inches(2.5))

    # 儲存
    document.save(output_word)


# 寫html-----------------------------------------------------------
def html(output_path, uuid_sample):
    output_txt = output_path + uuid_sample + '.txt'
    output_web = output_path + uuid_sample + '.html'

    # 讀檔
    article_file = open(output_txt, 'r', encoding='utf-8').read()
    Article = article_file.split()
    part = ""

    for i in range(len(Article)):
        div_head = "<div id = 'part{}'>".format(i + 1)
        paragraph = "<p id = 'paragraph{}'>{}\n</p>".format(i + 1, Article[i])
        picture_path = output_path + "image/" + uuid_sample + '_p{}.png'.format(i + 1)
        picture = "<img id='picture{}' src='{}' width='200'>".format(i + 1, picture_path)
        div_end = "</div>"
        part = part + div_head + paragraph + picture + div_end
    # 組合html
    html = """
    <html>
    <head>
    </head>
    <body>
    {}
    </body>
    </html>
    """.format(part)

    web = open(output_web, 'w', encoding='utf-8')
    web.write(html)


# ------------------------------------------------------------------------------

def pic(Output_path_cht, output_path_zh, uuid_sample, keyword, paragraph):
    option = webdriver.ChromeOptions()
    option.add_argument("--headless")
    driver = webdriver.Chrome(options=option)
    url = 'https://image.baidu.com/search/index?tn=baiduimage&ps=1&ct=201326592&lm=-1&cl=2&nc=1&ie=utf-8&word=' + keyword
    driver.get(url)

    driver.enconding = 'UTF-8'
    soup = BeautifulSoup(driver.page_source, 'html.parser')
    body = soup.find('div', attrs={'id': 'wrapper'})
    body = body.find('div', attrs={'id': 'imgContainer'})
    body = body.find('div', attrs={'id': 'imgid'})
    testpic_path = 'test_pic/'

    # testpic_path = '/home/ubuntu/content/gpt2-ml/scripts/test_pic/'

    i = 0
    count = 0
    for txt in body.find_all('div', attrs={'class': 'imgpage'}):
        txt = txt.find('ul', attrs={'class': 'imglist clearfix pageNum' + str(i)})
        i += 1
        for img in txt.find_all('li', attrs={'class': 'imgitem'}):
            img = img.find('img')
            img = img.attrs['data-imgurl']
            image = requests.get(img)
            print(image)
            # fp = open(testpic_path+  str(count+1) + '.png','wb')
            # fp.write(image.content)
            # fp.close()
            open(testpic_path + str(count + 1) + '.png', 'wb').write(image.content)
            count += 1
    print('爬取图片总数：', count)
    pathDir = os.listdir(testpic_path)
    s1 = random.sample(pathDir, paragraph)
    print(s1)
    i = 0
    for name in s1:
        i = i + 1
        shutil.copyfile(testpic_path + name, Output_path_cht + uuid_sample + "_p{}".format(i) + ".png")
    shutil.rmtree(testpic_path)
    os.mkdir(testpic_path)
    i = 0
    for num in range(int(paragraph)):
        i = i + 1
        shutil.copyfile(Output_path_cht + uuid_sample + "_p{}".format(i) + ".png",
                        output_path_zh + uuid_sample + "_p{}".format(i) + ".png")


# --------------------------------------------------------------------------------
def title_produce(output_path, uuid_sample):
    output_txt = output_path + uuid_sample + '.txt'
    output_file = output_path + uuid_sample + '_title.txt'

    title = codecs.open(output_txt, 'r', 'utf-8').read()
    tr4s = TextRank4Sentence()
    tr4s.analyze(text=title, lower=True, source="all_filters")

    for item in tr4s.get_key_sentences(num=1):
        print(item.sentence)
    f = open(output_file, 'w', encoding='utf-8')
    f.write(item.sentence)
    f.close()

# pic(Output_path_cht='scripts/test_pic/' , output_path_zh='scripts/zh_pic/', uuid_sample='123', keyword='test1',paragraph=1)
